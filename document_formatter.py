from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, List, Any
from io import BytesIO
import os
from datetime import datetime
from supabase import create_client

class ResumeFormatter:
    def __init__(self, margins: float = 0.8):
        self.doc = Document()
        self.margins = margins
        self._set_margins()
    
    def _set_margins(self) -> None:
        """Set document margins"""
        for section in self.doc.sections:
            section.top_margin = Inches(self.margins)
            section.bottom_margin = Inches(self.margins)
            section.left_margin = Inches(self.margins)
            section.right_margin = Inches(self.margins)
    
    def add_header(self, name: str, contact_info: str) -> None:
        """Add name and contact information"""
        name_para = self.doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(14)
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        contact = self.doc.add_paragraph()
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact.add_run(contact_info)
    
    def add_section(self, heading: str, content: List[Any], style: str = 'List Bullet') -> None:
        """Add a section with heading and content"""
        heading_para = self.doc.add_paragraph()
        heading_run = heading_para.add_run(heading)
        heading_run.bold = True
        heading_para.style = 'Heading 1'
        heading_format = heading_para.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)

        for item in content:
            if isinstance(item, tuple):
                p = self.doc.add_paragraph()
                p.add_run(f'{item[0]}: ').bold = True
                p.add_run(item[1])
            else:
                p = self.doc.add_paragraph(style=style)
                p.add_run(str(item))
    
    def create_resume(self, data: Dict[str, Any]) -> BytesIO:
        """Create resume from input data and return as BytesIO"""
        try:
            # Add header
            self.add_header(data['name'], data['contact_info'])
            
            # Add all sections
            for section in data['sections']:
                self.add_section(
                    section['heading'],
                    section['content'],
                    section.get('style', 'List Bullet')
                )
            
            # Save to BytesIO
            output = BytesIO()
            self.doc.save(output)
            output.seek(0)
            return output
            
        except Exception as e:
            raise Exception(f"Error creating resume: {str(e)}")

class DocumentFormatter:
    def __init__(self):
        # Initialize Supabase client
        self.supabase_url = os.getenv("SUPABASE_URL")
        self.supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
        if not self.supabase_url or not self.supabase_key:
            raise ValueError("Supabase environment variables are not properly set")
        self.supabase = create_client(self.supabase_url, self.supabase_key)
        self.resume_formatter = ResumeFormatter()

    async def process_and_save(self, file_content: bytes, original_filename: str) -> Dict[str, Any]:
        """Process the input file and save to Supabase storage"""
        try:
            # Load document
            doc = Document(BytesIO(file_content))
            
            # Extract data from the document
            data = self._extract_data_from_doc(doc)
            
            # Format the document using ResumeFormatter
            formatted_doc = self.resume_formatter.create_resume(data)

            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            sanitized_filename = ''.join(c for c in original_filename if c.isalnum() or c in '._- ')
            formatted_filename = f"formatted_{timestamp}_{sanitized_filename}"

            # Save to Supabase storage
            self.supabase.storage.from_("resume_templates").upload(
                formatted_filename,
                formatted_doc.getvalue()
            )

            return {
                "status": "success",
                "formatted_file_path": formatted_filename,
                "timestamp": timestamp
            }

        except Exception as e:
            print(f"Error in document formatting: {str(e)}")
            raise Exception(f"Failed to format document: {str(e)}")

    def _extract_data_from_doc(self, doc: Document) -> Dict[str, Any]:
        """Extract data from Document object into required format"""
        data = {
            "name": "Your Name",  # Default placeholder
            "contact_info": "your.email@example.com • Phone Number",  # Default placeholder
            "sections": []
        }

        current_section = None
        current_content = []

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading'):
                if current_section:
                    data["sections"].append({
                        "heading": current_section,
                        "content": current_content
                    })
                current_section = paragraph.text
                current_content = []
            elif paragraph.text.strip():
                current_content.append(paragraph.text)

        # Add the last section
        if current_section:
            data["sections"].append({
                "heading": current_section,
                "content": current_content
            })

        return data
