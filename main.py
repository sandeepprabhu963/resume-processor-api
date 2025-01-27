from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate
import google.generativeai as genai
import json
import os
from io import BytesIO
from typing import Dict, Any, List
from dotenv import load_dotenv
from datetime import datetime
import re
from supabase import create_client
import traceback
import time

load_dotenv()

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

# Initialize API clients
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("GEMINI_API_KEY environment variable is not set")
genai.configure(api_key=api_key)

supabase_url = os.getenv("SUPABASE_URL")
supabase_key = os.getenv("SUPABASE_SERVICE_ROLE_KEY")
if not supabase_url or not supabase_key:
    raise ValueError("Supabase environment variables are not properly set")
supabase = create_client(supabase_url, supabase_key)

def format_docx(doc: Document) -> Document:
    """Apply consistent formatting to the DOCX document."""
    try:
        print("Starting document formatting...")
        
        if not doc.paragraphs:
            raise ValueError("Document appears to be empty")
            
        # Set default font and spacing
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        
        # Format headings and content
        for paragraph in doc.paragraphs:
            if not paragraph.runs:
                continue
                
            # Detect and format section headers
            if (paragraph.text.isupper() or 
                paragraph.text.endswith(':') or 
                any(header in paragraph.text.lower() for header in 
                    ['summary', 'experience', 'education', 'skills', 'contact', 'objective'])):
                
                paragraph.style = doc.styles['Heading 1']
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(14)
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(4)
                
            # Format bullet points
            elif paragraph.text.strip().startswith('•'):
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        
        print("Document formatting completed successfully")
        return doc
        
    except Exception as e:
        print(f"Error in format_docx: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to format document: {str(e)}")

def docx_to_json(doc: Document) -> Dict[str, Any]:
    """Convert formatted DOCX document to JSON with enhanced structure."""
    try:
        print("Starting DOCX to JSON conversion...")
        
        if not doc.paragraphs:
            raise ValueError("Document appears to be empty")
        
        json_data = {
            "sections": {},
            "metadata": {
                "created_at": datetime.now().isoformat(),
                "format_version": "1.0"
            }
        }
        
        current_section = None
        section_content = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Enhanced section header detection
            if (text.isupper() or 
                text.endswith(':') or 
                any(header in text.lower() for header in 
                    ['summary', 'experience', 'education', 'skills', 'contact', 'objective'])):
                
                if current_section and section_content:
                    json_data["sections"][current_section] = {
                        "content": section_content,
                        "style": "regular"
                    }
                    section_content = []
                
                current_section = text.lower().replace(':', '').strip()
                
            else:
                if current_section:
                    para_data = {
                        "text": text,
                        "style": "bullet" if text.startswith('•') else "regular",
                        "formatting": {
                            "bold": any(run.bold for run in paragraph.runs),
                            "italic": any(run.italic for run in paragraph.runs)
                        }
                    }
                    section_content.append(para_data)
                elif text:  # Handle text before first section
                    current_section = "header"
                    section_content.append({
                        "text": text,
                        "style": "regular",
                        "formatting": {"bold": True}
                    })
        
        # Add the last section
        if current_section and section_content:
            json_data["sections"][current_section] = {
                "content": section_content,
                "style": "regular"
            }
        
        print("DOCX to JSON conversion completed successfully")
        return json_data
        
    except Exception as e:
        print(f"Error in docx_to_json: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to convert resume to JSON: {str(e)}")

def json_to_docx(template_doc: DocxTemplate, json_data: Dict[str, Any]) -> BytesIO:
    """Convert JSON back to DOCX format with enhanced formatting."""
    try:
        print("Starting JSON to DOCX conversion...")
        
        if not json_data.get("sections"):
            raise ValueError("JSON data contains no sections")
        
        # Create context for template
        context = {}
        
        for section_name, section_data in json_data.get("sections", {}).items():
            if not isinstance(section_data, dict) or "content" not in section_data:
                print(f"Warning: Invalid section data for {section_name}")
                continue
                
            # Convert section name to valid template variable
            template_var = re.sub(r'[^\w]', '_', section_name.lower())
            
            formatted_content = []
            for item in section_data.get("content", []):
                if not isinstance(item, dict) or "text" not in item:
                    continue
                    
                text = item.get("text", "").strip()
                if not text:
                    continue
                
                # Apply formatting
                if item.get("style") == "bullet" and not text.startswith('•'):
                    text = f"• {text}"
                
                # Handle special formatting
                if item.get("formatting", {}).get("bold"):
                    text = f"**{text}**"
                if item.get("formatting", {}).get("italic"):
                    text = f"*{text}*"
                
                formatted_content.append(text)
            
            context[template_var] = '\n'.join(formatted_content)
        
        # Render template
        template_doc.render(context)
        
        # Save to BytesIO
        output = BytesIO()
        template_doc.save(output)
        output.seek(0)
        
        print("JSON to DOCX conversion completed successfully")
        return output
        
    except Exception as e:
        print(f"Error in json_to_docx: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Failed to convert to DOCX: {str(e)}")

@app.post("/process-resume/")
async def process_resume(
    file: UploadFile = File(...),
    job_description: str = Form(...)
) -> StreamingResponse:
    try:
        if not file.filename.endswith('.docx'):
            raise HTTPException(status_code=400, detail="Invalid file format. Please upload a .docx file.")
            
        if not job_description.strip():
            raise HTTPException(status_code=400, detail="Job description cannot be empty")

        print(f"Processing resume: {file.filename}")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Read and process the document
        content = await file.read()
        if len(content) == 0:
            raise HTTPException(status_code=400, detail="Empty file uploaded")
        
        # Create Document objects and apply formatting
        doc = Document(BytesIO(content))
        doc = format_docx(doc)  # Apply consistent formatting
        template_doc = DocxTemplate(BytesIO(content))
        
        # Convert formatted DOCX to JSON
        json_data = docx_to_json(doc)
        
        # Save original JSON to Supabase
        filename = re.sub(r'[^\w\-\.]', '_', file.filename).strip('_').strip()
        original_json_path = f"original_{timestamp}_{filename}.json"
        
        try:
            json_str = json.dumps(json_data, ensure_ascii=False, indent=2)
            json_bytes = json_str.encode('utf-8')
            
            supabase.storage.from_("resume_templates").upload(
                original_json_path,
                json_bytes
            )
            print(f"Original JSON saved to: {original_json_path}")
        except Exception as e:
            print(f"Error saving original JSON: {str(e)}\n{traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"Failed to save original JSON: {str(e)}")
        
        # Optimize JSON using Gemini
        model = genai.GenerativeModel('gemini-pro')
        
        # Configure safety settings and generation config
        safety_settings = {
            "HARASSMENT": "block_none",
            "HATE_SPEECH": "block_none",
            "SEXUALLY_EXPLICIT": "block_none",
            "DANGEROUS_CONTENT": "block_none",
        }
        
        generation_config = {
            "temperature": 0.7,
            "top_p": 1,
            "top_k": 1,
            "max_output_tokens": 2048,
        }
        
        prompt = f"""
        Optimize this resume content for the following job description while maintaining the exact same structure:
        
        Resume Content:
        {json.dumps(json_data, ensure_ascii=False, indent=2)}
        
        Job Description:
        {job_description}
        
        Return ONLY the optimized JSON with the exact same structure.
        """
        
        # Implement retry mechanism for Gemini API
        max_retries = 3
        retry_delay = 1  # seconds
        
        for attempt in range(max_retries):
            try:
                response = model.generate_content(
                    prompt,
                    generation_config=generation_config,
                    safety_settings=safety_settings
                )
                
                if not response.text:
                    raise ValueError("Empty response from Gemini API")
                    
                response_text = response.text.strip()
                print(f"Gemini API response (attempt {attempt + 1}): {response_text[:200]}...")
                
                # Remove any markdown code block markers
                response_text = re.sub(r'^```json\s*|\s*```$', '', response_text)
                
                try:
                    optimized_json = json.loads(response_text)
                    if not isinstance(optimized_json, dict) or "sections" not in optimized_json:
                        raise ValueError("Invalid JSON structure in optimized content")
                    break  # Success, exit retry loop
                except json.JSONDecodeError as e:
                    print(f"Error parsing Gemini response: {str(e)}\nResponse text: {response_text}")
                    if attempt == max_retries - 1:  # Last attempt
                        raise HTTPException(status_code=500, detail="Failed to parse optimized content")
                    
            except Exception as e:
                print(f"Attempt {attempt + 1} failed: {str(e)}")
                if attempt == max_retries - 1:  # Last attempt
                    raise HTTPException(status_code=500, detail=f"Failed to optimize resume content: {str(e)}")
                time.sleep(retry_delay * (attempt + 1))  # Exponential backoff
        
        # Save optimized JSON
        optimized_json_path = f"optimized_{timestamp}_{filename}.json"
        optimized_json_str = json.dumps(optimized_json, ensure_ascii=False, indent=2)
        optimized_json_bytes = optimized_json_str.encode('utf-8')
        
        supabase.storage.from_("resume_templates").upload(
            optimized_json_path,
            optimized_json_bytes
        )
        print(f"Optimized JSON saved to: {optimized_json_path}")
        
        # Convert optimized JSON back to DOCX
        output = json_to_docx(template_doc, optimized_json)
        
        # Store optimization record
        try:
            supabase.table('resume_optimizations').insert({
                'original_resume_path': original_json_path,
                'optimized_resume_path': optimized_json_path,
                'job_description': job_description,
                'status': 'completed'
            }).execute()
            print("Optimization record stored in database")
        except Exception as e:
            print(f"Error storing optimization record: {str(e)}\n{traceback.format_exc()}")
        
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename=optimized_{filename}",
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
        
    except Exception as e:
        print(f"Processing error: {str(e)}\n{traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    return {"status": "healthy"}
